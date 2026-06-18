"""Minimal Markdown -> .docx converter for project docs.

Handles: ATX headings, fenced code blocks (shaded boxed), pipe tables,
blockquotes, bullet lists, horizontal rules, standalone images
(`![alt](path)` -> embedded picture + italic caption), and inline
**bold** / `code` / [text](url).

Not a general Markdown engine — scoped to this project's docs.
Usage:  python tools/md_to_docx.py <file.md> [out.docx]
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")
_IMG = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def add_inline(par, text):
    """Render inline bold / code / links into a paragraph."""
    for tok in _INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"; r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif tok.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", tok)
            par.add_run(m.group(1) if m else tok)
        else:
            par.add_run(tok)


def _shade_and_border(par, fill="F2F3F5", border="C9CDD2"):
    """Shaded background + thin box border (code-block look)."""
    pPr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single"); e.set(qn("w:sz"), "4")
        e.set(qn("w:space"), "6"); e.set(qn("w:color"), border)
        borders.append(e)
    pPr.append(borders)


def add_code_block(doc, lines):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Pt(6)
    par.paragraph_format.right_indent = Pt(6)
    par.paragraph_format.space_before = Pt(6)
    par.paragraph_format.space_after = Pt(8)
    _shade_and_border(par)
    r = par.add_run("\n".join(lines))
    r.font.name = "Consolas"; r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def add_image(doc, alt, path, base_dir):
    p = Path(path)
    if not p.is_absolute():
        p = (base_dir / p)
    if not p.exists():
        par = doc.add_paragraph()
        r = par.add_run(f"[missing image: {path}]"); r.italic = True
        return
    par = doc.add_paragraph(); par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    par.add_run().add_picture(str(p), width=Inches(6.3))
    if alt:
        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(alt); r.italic = True; r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def add_table(doc, rows):
    cols = len(rows[0])
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    for ri, cells in enumerate(rows):
        row = t.add_row().cells
        for ci, cell in enumerate(cells):
            row[ci].paragraphs[0].text = ""
            add_inline(row[ci].paragraphs[0], cell)
            if ri == 0:
                for run in row[ci].paragraphs[0].runs:
                    run.bold = True


def convert(md_path, docx_path):
    base_dir = Path(md_path).resolve().parent
    lines = Path(md_path).read_text(encoding="utf-8").splitlines()
    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    i, n = 0, len(lines)
    while i < n:
        line = lines[i]

        if line.lstrip().startswith("```"):
            i += 1; buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i]); i += 1
            add_code_block(doc, buf); i += 1; continue

        if line.strip().startswith("|") and i + 1 < n and re.match(
                r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            tbl = []
            while i < n and lines[i].strip().startswith("|"):
                if re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i]):
                    i += 1; continue
                cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                tbl.append(cells); i += 1
            add_table(doc, tbl); continue

        stripped = line.strip()

        m = _IMG.match(stripped)
        if m:
            add_image(doc, m.group(1), m.group(2), base_dir); i += 1; continue

        if stripped == "---":
            i += 1; continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3); i += 1; continue
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2); i += 1; continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1); i += 1; continue
        if stripped.startswith(">"):
            par = doc.add_paragraph(style="Intense Quote")
            add_inline(par, stripped.lstrip("> ").rstrip()); i += 1; continue
        if stripped.startswith("- "):
            par = doc.add_paragraph(style="List Bullet")
            add_inline(par, stripped[2:]); i += 1; continue
        if re.match(r"^\d+\.\s", stripped):
            par = doc.add_paragraph(style="List Number")
            add_inline(par, re.sub(r"^\d+\.\s", "", stripped)); i += 1; continue
        if not stripped:
            i += 1; continue

        buf = [stripped]; i += 1
        while i < n:
            nxt = lines[i].strip()
            if (not nxt or nxt.startswith(("#", "-", ">", "|", "```", "---", "!["))
                    or re.match(r"^\d+\.\s", nxt)):
                break
            buf.append(nxt); i += 1
        par = doc.add_paragraph()
        add_inline(par, " ".join(buf))

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "INSTALL.md"
    dst = sys.argv[2] if len(sys.argv) > 2 else str(Path(src).with_suffix(".docx"))
    convert(src, dst)